#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
DOCX格式输出
"""

import os
from typing import Dict, Any, List
from pathlib import Path
import datetime

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..models import Document, Paragraph, Chapter
from ..utils import ensure_dir

def format_as_docx(document: Document, options: Dict[str, Any], output_path: str) -> str:
    """
    将文档导出为DOCX格式
    
    Args:
        document: 文档对象
        options: 格式选项
        output_path: 输出路径
        
    Returns:
        输出文件路径
    """
    # 解析选项
    bilingual = options.get("bilingual_output", False)
    include_titles = options.get("include_titles", True)
    author = options.get("author", "Unknown")
    
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    ensure_dir(output_dir)
    
    # 创建DOCX文档
    docx = DocxDocument()
    
    # 设置文档属性
    docx.core_properties.title = document.title
    docx.core_properties.author = author
    docx.core_properties.language = document.target_language if not bilingual else document.source_language
    
    # 添加标题
    title = docx.add_heading(document.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加元信息段落
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    info_para = docx.add_paragraph()
    info_para.add_run(f"源语言: {document.source_language}").italic = True
    info_para.add_run(f" | 目标语言: {document.target_language}").italic = True
    info_para.add_run(f" | 生成日期: {today}").italic = True
    docx.add_paragraph()  # 空行
    
    # 设置样式
    style = docx.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # 按章节添加内容
    if document.chapters:
        # 有章节结构的情况
        for chapter_id, chapter in sorted(document.chapters.items(), key=lambda x: x[0]):
            # 添加章节标题
            heading = docx.add_heading(chapter.title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 章节标题段落
            title_para = None
            for para_id in chapter.paragraphs:
                para = document.paragraphs.get(para_id)
                if para and para.is_title:
                    title_para = para
                    break
            
            # 如果找到标题段落且为双语模式，添加译文标题
            if title_para and bilingual and title_para.is_translated:
                translated_heading = docx.add_heading(title_para.translated, level=2)
                translated_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 设置标题样式为斜体
                for run in translated_heading.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)
            
            # 添加章节内容段落
            for para_id in chapter.paragraphs:
                para = document.paragraphs.get(para_id)
                if para and not para.is_title:  # 跳过标题段落，已单独处理
                    # 添加原文
                    p = docx.add_paragraph(para.content)
                    
                    # 如果是双语模式，添加译文
                    if bilingual and para.is_translated:
                        trans_p = docx.add_paragraph(para.translated)
                        trans_p.style = 'Quote'
                        for run in trans_p.runs:
                            run.italic = True
                            run.font.color.rgb = RGBColor(80, 80, 80)
            
            # 每章结束添加分页符
            docx.add_page_break()
    else:
        # 无章节结构，将所有内容作为普通段落添加
        for para_id, para in sorted(document.paragraphs.items(), key=lambda x: x[0]):
            if not para.is_title or include_titles:
                # 添加原文
                p = docx.add_paragraph(para.content)
                
                # 如果是双语模式，添加译文
                if bilingual and para.is_translated:
                    trans_p = docx.add_paragraph(para.translated)
                    trans_p.style = 'Quote'
                    for run in trans_p.runs:
                        run.italic = True
                        run.font.color.rgb = RGBColor(80, 80, 80)
    
    # 保存文档
    docx.save(output_path)
    
    # 如果是双语模式，再生成两个单语言的版本
    if bilingual:
        # 源语言版本
        source_path = f"{os.path.splitext(output_path)[0]}_source.docx"
        format_single_language_docx(document, False, options, source_path)
        
        # 目标语言版本
        target_path = f"{os.path.splitext(output_path)[0]}_target.docx"
        format_single_language_docx(document, True, options, target_path)
    
    return output_path

def format_single_language_docx(document: Document, target_language: bool, options: Dict[str, Any], output_path: str) -> str:
    """
    将文档导出为单语言DOCX格式
    
    Args:
        document: 文档对象
        target_language: 是否为目标语言
        options: 格式选项
        output_path: 输出路径
        
    Returns:
        输出文件路径
    """
    # 确保输出目录存在
    output_dir = os.path.dirname(output_path)
    ensure_dir(output_dir)
    
    # 解析选项
    include_titles = options.get("include_titles", True)
    author = options.get("author", "Unknown")
    
    # 创建DOCX文档
    docx = DocxDocument()
    
    # 设置文档属性
    docx.core_properties.title = document.title
    docx.core_properties.author = author
    docx.core_properties.language = document.target_language if target_language else document.source_language
    
    # 添加标题
    title = docx.add_heading(document.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加元信息段落
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    lang = document.target_language if target_language else document.source_language
    info_para = docx.add_paragraph()
    info_para.add_run(f"语言: {lang}").italic = True
    info_para.add_run(f" | 生成日期: {today}").italic = True
    docx.add_paragraph()  # 空行
    
    # 设置样式
    style = docx.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # 按章节添加内容
    if document.chapters:
        # 有章节结构的情况
        for chapter_id, chapter in sorted(document.chapters.items(), key=lambda x: x[0]):
            # 获取章节标题段落
            title_para = None
            for para_id in chapter.paragraphs:
                para = document.paragraphs.get(para_id)
                if para and para.is_title:
                    title_para = para
                    break
            
            # 添加章节标题
            if title_para:
                title_text = title_para.translated if target_language and title_para.is_translated else title_para.content
                heading = docx.add_heading(title_text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加章节内容段落
            for para_id in chapter.paragraphs:
                para = document.paragraphs.get(para_id)
                if para and not para.is_title:  # 跳过标题段落，已单独处理
                    # 根据语言选择内容
                    content = para.translated if target_language and para.is_translated else para.content
                    p = docx.add_paragraph(content)
            
            # 每章结束添加分页符
            docx.add_page_break()
    else:
        # 无章节结构，将所有内容作为普通段落添加
        for para_id, para in sorted(document.paragraphs.items(), key=lambda x: x[0]):
            if not para.is_title or include_titles:
                # 根据语言选择内容
                content = para.translated if target_language and para.is_translated else para.content
                p = docx.add_paragraph(content)
    
    # 保存文档
    docx.save(output_path)
    
    return output_path 